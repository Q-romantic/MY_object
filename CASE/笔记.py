'''替换
# 方法一：
(.*?): (.*)  或  (.+): (.+)  或  (.): (.)$
'$1': '$2',
# 方法二：
s = """Host: sou.zhaopin.com
origin: https://sou.zhaopin.com
referer: https://sou.zhaopin.com/?p=3&jl=765&kw=python&kt=3"""

def str_to_dic(s: str = '', flags: [bool, int] = 0) -> dict:
    import re
    # pattern = '^(.*?):\s{0,}(.*)$'    # 上下等价
    pattern = '^(.*?):\s*(.*)$'
    h = headers
    for line in s.splitlines():
        if line == '' or line.startswith('#'):
            continue
        s = re.sub(pattern, '{\'\\1\': \'\\2\'}', line)
        h.update(eval(s))
        if flags:
            print(re.sub(pattern, "\'\\1\': \'\\2\',", line))
    return h
# 方法三：
headers = dict([line.split(': ' if " " in line else ':', 1) for line in s.split("\n") if line != '' and "#" not in line])
'''

'''0、获取网页数据
import pprint

pprint.pprint(response.json())  # python美化打印的标准库
books = ['云游', '人生十二法则', '一场游戏一次消遣']
books_str = pprint.pformat(books)  # 不直接打印，而是返回一个字符串。这个字符串在语法上是完全符合 Python 代码规范的。

1）
import requests
import json
import filetype

p_ip = requests.get('https://proxypool.scrape.center/random').text
proxies = {'http': f'http://{p_ip}', 'https': f'http://{p_ip}'}
data = json.dumps(data)  # Requests发Post请求data里面嵌套字典
filename = 'C:/Users/11390/Desktop/Snipaste_2022-04-02_10-03-56.png'
mime_type = filetype.guess(filename).mime  # 获取文件类型
files = {  # 适用于单文件提交
    # 'files[]': ['1.png', f'filename={filename}', 'image/png'],  # 上下两种方式传入均可
    'files[]': ['1.png', filename, mime_type],
}
response = requests.post(url, headers=headers, params=params, data=data, json=data, files=files, proxies=proxies)
print(response.text)

2）
from requests_html import HTMLSession
session = HTMLSession()

response = session.get(url, headers=headers, params=data)
print(response.text)

'''
###############################################################################
'''
1、本地xml转网页访问，xpath取值
from lxml import etree
import os
import requests
from requests_testadapter import Resp

path = r'C:/Users/11390/Desktop/EMS-info/logs/SGA00802W1.xml'


class LocalFileAdapter(requests.adapters.HTTPAdapter):
    def build_response_from_file(self, request):
        file_path = request.url[7:]
        with open(file_path, 'rb') as file:
            buff = bytearray(os.path.getsize(file_path))
            file.readinto(buff)
            resp = Resp(buff)
            r = self.build_response(request, resp)

            return r

    def send(self, request, stream=False, timeout=None,
             verify=True, cert=None, proxies=None):
        return self.build_response_from_file(request)


requests_session = requests.session()
requests_session.mount('file://', LocalFileAdapter())
response = requests_session.get('file://' + path)

xml = etree.HTML(response.text)
print(xml.xpath('//storagecell/object[1]/objectname/text()'))
'''
###############################################################################
'''
2.0、网页xpath或css选择器取值(推荐)
import parsel

selector = parsel.Selector(response.text)
title = selector.css('.bookname h1::text').get()  # css 语法相对路径取文本内容
title_1 = selector.xpath('//[@class="bookname"]/h1/text()').get()
content = ''.join(selector.css('#content::text').getall())
'''
###############################################################################
'''
2.1、网页xpath取值
from lxml import etree

path = r'C:/Users/11390/Desktop/EMS-info/logs/SGA00802W1.xml'

with open(path, mode="r", encoding="utf-8") as f:
    r = f.read()
    xml = etree.HTML(r)
    # xml = etree.HTML(r)    # 三种方法
    # xml = etree.parse(r)
    # xml = etree.XML(r)
    print(xml.xpath('//storagecell/object[1]/objectname/text()'))
'''
###############################################################################
'''
2.2、网页xpath取值
response = requests.get(url)
xml = etree.HTML(response.text)
link_list = xml.xpath('//*[@id="mp-editor"]/p/img')
for i in link_list:
    link_url = i.attrib['src']
'''
###############################################################################
'''
2.3、网页标签取值
from bs4 import BeautifulSoup

path = r'C:/Users/11390/Desktop/EMS-info/logs/SGA10802DX.xml'

with open(path, "r", encoding="utf-8") as f:
    html = f.read()
    soup = BeautifulSoup(html, features="lxml")  # 方法1
    # soup = BeautifulSoup(html, "html.parser") # 方法2
    # print(soup.storagecell.object.objectname)
    print(soup.objectname.text)
    # print(soup.html.find_all('object')[0].objectname)
'''
###############################################################################
'''
4、网页返回json数据取值
from jsonpath import jsonpath  # pip install jsonpath

response_json = requests.get(url, headers=headers, params=data).json()
song_name = jsonpath(response_json, '$..name')  # 歌曲名
'''
###############################################################################
'''
1、python文件命名得以字母或下划线开头才可以调用
调用方法：
from 目录.文件名 import 类名或函数名
from 文件名 import 类名或函数名

1.1、若调用类方法用
'类名().函数名()'
1.2、若调用函数用
'函数名()'
2、装饰器按类编写，调用时使用
'@类名().方法函数名'，其余用
'@函数名'
3、在类中方法函数的第一个参数可以时任意字符或下划线组合，并且不参与(实参、形参)参数传递
使用参考说明：

class Test(object):
    age = 18

    def funtion():
        print('无参数function')

    def normal_funtion(any):
        print('normal_function')

    @classmethod
    def class_method(cls):
        print(cls.__name__)

    @staticmethod
    def static_method():
        # print(Test.age)
        print('static_method')


any = Test()
# any.funtion()             # 无法调用，因为通过实例调用时，实例会被当作第一个参数传给normal_funcion。
any.normal_funtion()  # 可以调用
any.class_method()  # 可以调用，当@classmethod检测到是通过实例调用时，会把当前实例的__class__，当作cls传递，可以访问类属性，但无法访问实例属性
any.static_method()  # 可以调用，当@staticmethod检测到是通过实例调用时，会在当前实例的类中调用静态方法，无法访问类或实例的属性，但是可以访问全局变量(比如全局变量Test类)

Test.funtion()  # 可以调用，类调用时不传递参数，刚好function也不需要参数，无法访问类或实例的属性
Test().normal_funtion()  # 可以调用，类调用时需要传递参数
Test.class_method()  # 可以调用，类方法，通过类调用时，类会被当作cls传递给函数，可以访问类属性,无法访问实例属性
Test.static_method()  # 可以调用，静态方法，等于定义在类内的函数，无法访问类或实例的属性，但是可以访问全局变量(比如全局变量Test类)
'''
###############################################################################
'''
生成二维码 + 自定义背景图
from MyQR import myqr

myqr.run(
    words='https://u.wechat.com/MCl10WdUUKVBp0K3JnYlPd4',
    # 扫描二维码后，显示的内容，或是跳转的链接
    version=5,  # 设置容错率
    level='H',  # 控制纠错水平，范围是L、M、Q、H，从左到右依次升高
    picture=r'C:\\Users\\11390\Desktop\\444.png',  # 图片所在目录，可以是动图
    colorized=True,  # 黑白(False)还是彩色(True)
    contrast=1.0,  # 用以调节图片的对比度，1.0 表示原始图片。默认为1.0。
    brightness=1.0,  # 用来调节图片的亮度，用法同上。
    save_name='J9360.png',  # 控制输出文件名，格式可以是 .jpg， .png ，.bmp ，.gif
    # save_dir=r'C:\\Users', # 图片存储位置
)
'''
###############################################################################
'''
根据字典dict的值value查找key
1、方法一
def get_dict_key(dic, value):
    key = list(dic.keys())[list(dic.values()).index(value)]
    return key

2、方法二
def get_dict_key(dic, value):
    keys = list(dic.keys())
    values = list(dic.values())
    idx = values.index(value)
    key = keys[idx]
    return key
'''
###############################################################################
'''
图片按像素分解为RGB颜色数值代码以便查看或替换颜色
from PIL import Image

img = Image.open(r"C:\\Users\\11390\\Desktop\\20161020052802241.jpg")  # 读取系统的内照片
print(img.size)  # 打印图片大小
print(img.getpixel((0, 0)))
print(img.getpixel((img.size[0] - 1, img.size[1] - 1)))

width = img.size[0]  # 横向长度X→
height = img.size[1]  # 纵向高度Y↓
for i in range(0, width):  # 遍历所有长(高)度的点
    for j in range(0, height):  # 遍历所有宽度的点
        data = img.getpixel((i, j))  # 打印该图片的所有点
        # print(data, end='')#打印每个像素点的颜色RGBA的值(r,g,b)
        if (i <= 100 and j <= 200) or (i >= 1142 and j >= 2108):  # RGBA的r值大于170，并且g值大于170,并且b值大于170
            img.putpixel((i, j), (255, 0, 0))  # 则这些像素点的颜色改成大红色
img = img.convert("RGB")  # 把图片强制转成RGB
img.save(r"C:\\Users\\11390\\Desktop\\455.png")  # 保存修改像素点后的图片
'''
###############################################################################
'''
练习：日期计算星座


def borth_to_consteliation(month, day):
    consteliation = ('摩羯座', '水瓶座', '双鱼座', '白羊座', '金牛座', '双子座', '巨蟹座', '狮子座', '处女座', '天秤座', '天蝎座', '射手座', '摩羯座')
    # 定义星座
    date = (20, 19, 20, 20, 21, 21, 22, 23, 23, 23, 22, 21)  # 定义日期

    if day <= date[month - 1]:
        return consteliation[month - 1]
    else:
        return consteliation[month]


while 1:
    month = int(input("请输入月份: "))
    day = int(input("请输入日期: "))
    info = str(month) + "月" + str(day) + "日" + "对应的星座是:" + borth_to_consteliation(month, day)
    print()
    print(info)
'''
###############################################################################
'''
路径使用通配符
import glob

path = "C:\\Users\\11390\\Desktop\\"
li = glob.glob(path + "\*.jpg")
'''
###############################################################################
'''
列表区间取值问题
nums = [1, 3, -1, -3, 5, 3, 6, 7]
k, i = 3, 0

while i <= len(nums) - k:
    # print(nums[i:i + k])
    print(max(nums[i:i + k]))
    i += 1
'''
###############################################################################
'''
命令行选项、参数和子命令解析器(cmd终端参数分解传参)
# 用法1：python39 豆瓣电影250（正则+csv格式写入）.py 1,2,3 10
import sys

gpus = sys.argv[1]
# gpus = [int(gpus.split(','))]
batch_size = sys.argv[2]
print(gpus)
print(batch_size)

for line in sys.stdin:  # 循环等待输入
    n = eval(line)

# 用法2：python39 豆瓣电影250（正则+csv格式写入）.py --gpus=0,1,2 --batch_size=10
import tensorflow._api.v2.compat.v1 as tf

tf.app.flags.DEFINE_string('gpus', None, 'gpus to use')
tf.app.flags.DEFINE_integer('batch_size', 5, 'batch size')

FLAGS = tf.app.flags.FLAGS

print(FLAGS.gpus)
print(FLAGS.batch_size)

# 用法3：python39 豆瓣电影250（正则+csv格式写入）.py --seed=100
import argparse

parser = argparse.ArgumentParser(description='manual to this script')

parser.add_argument('--sparse', action='store_true', default=False, help='GAT with sparse version or not.')
parser.add_argument('--seed', type=int, default=72, help='Random seed.')
parser.add_argument('--epochs', type=int, default=10000, help='Number of epochs to train.')

args = parser.parse_args()
argss = vars(parser.parse_args())   # 可转为字典

# print(args)
print(args.sparse)
print(args.seed)
print(args.epochs)
'''
###############################################################################
'''等长列表转字典（一一对应）
方法一：
key_list = ['a', 'b', 'c']
value_list = [1, 2, 3]
dic = dict(zip(key_list, value_list))
# {'a': 1, 'b': 2, 'c': 3}

方法二：
dic = {key_list[i]: value_list[i] for i in range(len(key_list))}
'''
###############################################################################
'''生成伪数据
from faker import Faker
import pandas as pd

fake = Faker(["zh_CN"])
Faker.seed(0)


def get_data():
    key_list = ["姓名", "详细地址", "所在省份", "手机号", "身份证号", "出生年月", "邮箱"]
    name = fake.name()
    address = fake.address()
    province = address[:3]
    number = fake.phone_number()
    id_card = fake.ssn()
    birth_date = id_card[6:14]
    email = fake.email()
    info_list = [name, address, province, number, id_card, birth_date, email]
    person_info = dict(zip(key_list, info_list))
    return person_info


df = pd.DataFrame(columns=["姓名", "详细地址", "所在省份", "手机号", "身份证号", "出生年月", "邮箱"])
for i in range(1):
    person_info = [get_data()]
    df1 = pd.DataFrame(person_info)
    df = pd.concat([df, df1])
# df.to_excel("模拟数据.xlsx", index=None)
'''
###############################################################################
'''
列表形式写入csv文件
import csv

with open('data.csv', mode='w', encoding='utf-8', newline='') as f:
    csvwriter = csv.writer(f)
    TT = ['prodid', '一级分类', '二级分类', '品名', '最低价', '平均价', '最高价', '规格', '产地', '单位', '发布日期']
    csvwriter.writerow(TT)
'''
###############################################################################
'''
解决网页中文编码乱码
url = 'https://www.baidu.com/'

# 方法一：
resp = requests.get(url).content.decode()
# 方法二：
resp = requests.get(url)
resp.encoding = 'gbk'  # 默认写入 txt 以 gbk 方式
resp.encoding = 'utf-8'
resp.encoding = 'utf-8-sig'  # 一般写入txt或csv报错才使用，UnicodeEncodeError: 'gbk' codec can't encode character '\uc0ac' in position 14: illegal multibyte sequence
resp.encoding = 'unicode_escape'  # 将unicode的内存编码值进行存储，读取文件时在反向转换回来，例如："\u002F\u002F"
resp.encoding = resp.apparent_encoding  # 从网页的内容中分析网页编码
# print resp.text.encode("latin1").decode("utf-8")  # 这里的utf-8改成gbk就会报错，但有时候又情况相反，所以觉得这不是个正规办法
print(resp.text)

import io
import sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='gb18030')  # 改变标准输出的默认编码, 未验证

s = '我'
print(ord(s))  # 解码：chr(s)  # 转 ASCII

from urllib.parse import quote
print(quote(s, 'utf-8'))  # 解码：unquote(s, 'utf-8')

print(s.encode())  # 解码：s.decode()
print(s.encode('utf-8'))
print(s.encode('utf-16'))
print(s.encode('UTF-16BE'))
print(s.encode('UTF-16LE'))
print(s.encode('utf-8-sig'))
print(s.encode('gbk'))
print(s.encode('Shift_JIS'))
print(s.encode('BIG5'))
print(s.encode('GB2312'))
print(s.encode('GB18030'))
print(s.encode('ANSI'))
print(s.encode('unicode_escape'))
'''
###############################################################################
'''
# 视频分片或文件合并
# MAC / Linux：
# cat 1.ts 2.ts 3.ts > 123.mp4
# windows：
# copy /b 1.ts + 2.ts + 3.ts 123.mp4
# copy /b *.ts 123.mp4
# 还可以将所有ts文件放入一个文件夹，然后有压缩工具以存储方式压缩
'''
###############################################################################
'''
可迭代对象前加“ * ”号含义

def add(a, b):
    return a + b

data = {'a': 4, 'b': 3}
i = add(*data)  # 取key
j = add(**data)  # 取value
print(i)
print(j)
data = [4, 3]
i = add(*data)
print(i)
'''
###############################################################################
'''
reres插件映射本地js文件，结合jQuery搜索页面内容位置
// 理想国残书已搜到
$(function(){
if ($(".title").text()
indexof("理想国残书已收到") > -1){
    alert("存在了");
} else {
    window.location.href =$(".next a:eq(0)").attr("href");
}
})
'''
###############################################################################
'''生成美观的ASCII格式的表格
import prettytable as pt
tb = pt.PrettyTable()
tb.field_names = ['电影名称', 'Img', '类型', 'other', '上映时间', 'detail', '评分']
tb.add_row(li)
print(tb)
'''
###############################################################################
'''
保存数据
1) 保存文本内容
with open('./html.txt', 'a', encoding='utf-8') as f:
    f.write(response.text + '\n')

2) 保存图片、视频、音频等二进制内容
with open('./1.mp3', 'wb') as f:
    f.write(response.content)

3) 特殊图片
src = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAABgAAAAYCAYAAADgdz34AAAAAXNSR0IArs4c6QAABFhJREFUSEttlm2IVGUUx3/nzuzs+kaUTYHE0qv5Umwv+CVDe9MKgyAyqFxlKwSpKNN9mdnVvbo6M+22KpQQLG7hWn3YIIKKdEGz2CKCaiEVjQxCpNy0om3VbeaeeJ773LkzTs+Xucw9z/mf/3n5nytUHRV8BF8CVmqC61mAx8N43E/AIoTLAAHOoBxD+JSAAzQwZu/46uGjIBq5NcbxsQYS0Ko3kKIb5TFgBphLlIB/3XMD4LmL/wAfkqCXHhmzgQ2LsbUnBoicZ3QNQh+QRjmDx8coB/EYJck4EygejSRpQnkE4UHgSuAsSjt52RMykSAGiJ1vRtjiwN9E2UleToBKSNv82mvuWZQ2vZkE7Qgt7t5mctITgUiZUlafAwZcGjLkpN86XEuSOZSiiMr5NEGZNPkmdaJkdQNQAJJACzl52/gOI8rqHcB+S1XYwHbZYSM4gth8+jqdi6wBlgJFYIR6hvFl0ga4ELUBdOorKP02tUXupk9+DAEyug/haWCQnDxrnZtjLq3XK5jG+8C9gM2rK/AhzvM4O+VclX1W3wCeRxhiu6wWNukiSnwOnKfIYnrlWFVUWe0EtgHHUfoQTNd0A/MI6KcgG20qfRL4UiSjCxC+tEEELBWy2gr0ErCPgjSHNTGRivKi1jOLk8AclLvIy1c2/i69h4BDKKeY4EZel4tVIFndAzwDtBmAz2xuA56iIO9V9XEIcAqlRD2NnEb5g4BrSDGNcxYsxWxbC8NiJZ6tWUabEfaaWgmd+itKmiK32PRUtaTtjtuso5x8X36X0TaEV1FGqWdJ3GGunbv0OgJOAH8bBmZKJ0mRLkdSMerWuaogooSOzXTfCkwn4AkKMhyzdgC+zmSK34xNJcDV+DIRM3D9Eo1+uzaS4DBwLXDacCInu8tzEUYSDqRp6ynGIwCTy1lMMZfX5OcagGjKQ9pfAA0oK8jL15fqTvlum84nyQ8o40JGDyAsQ1lNXoZqLkUAHXo5Ht/aWZjkTnbJn5WaYwlEbDv0STzeRTkct2k8ZEnbz1XHUW/XJlIU6ZEjNUwrAbI6BKwK27RDb8ezg1ZyvX60moVzvlFnUMc6BGGSAcug3HEV0Wd0HsKoa4IlkVTsRWgGdpOTF6qo+xoy6tIHCBhxxJaTkxGid1ZWnER36iBKC8o75GVVCNCqN1FnUdNlsTP5NCcSsrAGu1DM/y+Rl7Pl4TJ2ZsCyuh7YYXdDwDIK8l0s1xltQRh0atlRlmuzQq2qOvmI2tFM7UKzXp1ch0pq5LoOZS15GYjlOu6UTXhstT6Et+y0bpPj/1vQKP8ZnYvwMrAuxKabvGyNF07ULRFIuHhMJLOB31H2I3xEkTECfmEmwgRpUixGuI+AFQhXWVtotYumZmVGIPHUNuFhZOFRt/SNhdkFF9werwNbC1NDs6U/oMgW+uSnS2ej+qsi1PXws8VEcQEDtBzlIYT5YCNVlL/w+IYSB1E+4SRH3ear+Wz5D7DBJjcA4p3NAAAAAElFTkSuQmCC"
# 1、信息提取
result = re.search("data:image/(?P<ext>.*?);base64,(?P<data>.*)", src, re.DOTALL)
ext = result.groupdict().get("ext")
data = result.groupdict().get("data")
# 2、base64解码(如下2种方法)
img = base64.urlsafe_b64decode(data)
img = base64.b64decode(data)
# 3、二进制文件保存
# filename = "{}.{}".format(uuid.uuid4(), ext)
# with open(filename, "wb") as f:
#     f.write(img)

注意：Windows文件名不能包含如下字符
import re

str = re.sub('[\/:*?"<>|]', '-', str)  # []过滤里面任意字符
'''
###############################################################################
''' # PIL转: numpy数组 <---> 图片Path <---> PIL <---> 二进制 <---> numpy数组
参考：https://blog.csdn.net/weixin_37763340/article/details/116717743
from PIL import Image
from io import BytesIO
import numpy as np
import cv2

image = cv2.imread('C:/Users/11390/Desktop/1.png')
cv2.imshow('title',image)
cv2.waitKey(0)
cv2.destroyAllWindows()
print(image, type(image))  # <class 'numpy.ndarray'>

img = Image.open('C:/Users/11390/Desktop/1.png')
img.show()
print(img, type(img))  # <class 'PIL.PngImagePlugin.PngImageFile'>

bytesIO = BytesIO() # 注意：针对每个图片都要新建，这里并不代表赋值，感觉像是新建一个实例化对象
img.save(bytesIO, format='PNG')  # format='JPEG'，实际未保存
data = bytesIO.getvalue()
print(data, type(data))  # <class 'bytes'>

data = np.asarray(bytearray(data), dtype="uint8")
image = cv2.imdecode(data, cv2.IMREAD_COLOR)
print(image, type(image))  # <class 'numpy.ndarray'>

# numpy矩阵转二进制（数组转二进制）: image ---> data
data = cv2.imencode('.jpg', image)[1].tobytes()
print(data, type(data))  # <class 'bytes'>

# 二进制转PIL: data --> img
img = Image.open(BytesIO(resp.content))
print(img)
img.show()

# 二进制转numpy矩阵（二进制转数组）: data ---> image
data = np.asarray(bytearray(resp.content), dtype="uint8")
image = cv2.imdecode(data, cv2.IMREAD_COLOR)
print(image)
cv2.imshow('title',image)

# 保存本地
cv2.imwrite('C:/Users/11390/Desktop/1122.png', image)
img.save('C:/Users/11390/Desktop/1133.png')
'''
###############################################################################
'''json.loads 和 eval 函数区别：
s = '{"one":1} '
print(eval(s), json.loads(s))
s = "{'one':1}"
print(eval(s), json.loads(s))   # json不认单引号，json中的字符串需要用双引号包起来
'''
###############################################################################
''' 读写入csv文件，参考：https://zhuanlan.zhihu.com/p/38537335
import csv

f = open('ttt.csv', mode='wt', encoding='utf-8', newline='')
# 方式一（字典索引写入）
headers = [
    'a1',
    '表头2',
    'a3'
]
dit_writer = csv.DictWriter(f, fieldnames=headers)
dit_writer.writeheader()  # 写入表头
dit = {
    'a1': 1,
    'a3': 2
}
dit_writer.writerow(dit)  # 按表头字典写入

# 方式二（按行顺序写入）
writer = csv.writer(f)
writer.writerow(('hello', 'world'))  # 单行写入，传入数据的格式为列表元组格式
writer.writerows([('hello', 'world'), ('I', 'love', 'you')])  # 多行行写入，传入数据的格式为列表元组格式
f.close()

# 方式一（按行顺序读取）
f = open('ttt.csv', mode='rt', encoding='utf-8', newline='')
reader = csv.reader(f, delimiter=',', quotechar='|')
# delimiter 指定分隔符，默认为逗号
# quotechar 表示引用符，未能理解？？？
for row in reader:
    # print(', '.join(row))
    print(row)
f.close()

# 方式二（字典索引读取）
f = open('ttt.csv', mode='rt', encoding='utf-8', newline='')
reader = csv.DictReader(f, delimiter=',')
for row in reader:
    print(row['a1'], row['a3'])
f.close()


import pandas as pd
# df = pd.DataFrame([['小明', '14'], ['小刚', '15']], columns=headers)
df = pd.DataFrame([['小明', '14'], ['小刚', '15']])
df.to_csv('ttt.csv', mode='at', index=False, sep=',')
'''
###############################################################################
'''读写xls文件
import xlrd
from xlutils.copy import copy

data = xlrd.open_workbook('1.xls')  # 打开xls文件
table = data.sheets()[0]
rows_old = table.nrows  # 获取表的行数
new_workbook = copy(data)
new_worksheet = new_workbook.get_sheet(0)
for i in range(0, rows_old):
    profile_url = weiboname_to_url(table.row_values(i)[0])    # 读取
    for j in range(1, 2):
        new_worksheet.write(i, j, profile_url)  # 从第一行开始写
new_workbook.save('2.xls')  # 保存工作簿
'''
###############################################################################
'''Python替换字符串中的反斜杠\
s = 'cdp\nd'
result = eval(repr(s).replace('\\', '@'))
print(result)
'''
###############################################################################
'''scrapy框架
python39 -m pip install scrapy

scrapy startproject chinadaily  # 创建爬虫项目
cd .\chinadaily
scrapy genspider 爬虫文件名 网站域名     # 创建爬虫文件
scrapy crawl english    # 运行

'''
###############################################################################
'''设置有效期限，到期失效
代码加密参考：https://zhuanlan.zhihu.com/p/54296517
import time
 
def now():
    return time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time()))
     
s = '2022-10-26 00:00:00'
print(now())
if now() > s:
    print( '已过期,无法正常使用')
else:
    print('未过期,可以正常使用')

#s指的是你想要的失效时间。
#now()指的是当前时间。
'''
###############################################################################
''' json 转 csv/excl
import pandas as pd
import csv
import json

data = """
{
    "Results":
        [
            {"id": "1", "Name": "Jay"},
            {"id": "2", "Name": "Mark"},
            {"id": "3", "Name": "Jack"}
        ],
    "status": ["ok"]
}
"""
info = json.loads(data)['Results']
# 方法一：
df = pd.json_normalize(info)
df.to_csv("samplecsv.csv", index=False)

# # 方法二：
print(info[0].keys())
with open("samplecsv.csv", 'w', newline='') as f:
    wr = csv.DictWriter(f, fieldnames=info[0].keys())
    wr.writeheader()
    wr.writerows(info)
'''
###############################################################################
''' xml 转 json
import xmltodict
xml = """
<Mashbits>
  <setting x="5" y="5" star="12"></setting>
  <object x="1" y="4" color="GRAY" isColored="true" isBlocked="true" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">NONE</object>
  <object x="3" y="4" color="GRAY" isColored="true" isBlocked="true" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">NONE</object>
  <object x="2" y="2" color="BLUE" isColored="true" isBlocked="false" isFreeze="false" isInverted="true" isCrashed="false" movement="NONE" id="-1">NONE</object>
  <object x="0" y="0" color="BLUE" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">RB</object>
  <object x="4" y="0" color="PINK" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">RP</object>
  <object x="4" y="4" color="BLUE" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">MB</object>
  <object x="0" y="4" color="PINK" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">MP</object>
  <object x="3" y="0" color="GRAY" isColored="true" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="VERTICAL" id="1">NONE</object>
  <object x="0" y="2" color="GRAY" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">OB</object>
  <object x="4" y="2" color="GRAY" isColored="false" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="NONE" id="-1">OB</object>
  <object x="1" y="0" color="GRAY" isColored="true" isBlocked="false" isFreeze="false" isInverted="false" isCrashed="false" movement="VERTICAL" id="1">NONE</object>
</Mashbits>
"""
xmlparse = xmltodict.parse(xml, encoding='utf-8')
# pprint.pprint(xmlparse)
jsonstr = json.dumps(xmlparse, indent=4).replace('@', '').replace('#', '')
jsonstr = json.loads(jsonstr)
pprint.pprint(jsonstr)
'''
###############################################################################
'''Python3的URL编码解码，呈现的结果是 --> %xx%xx
from urllib.parse import quote
text = quote(text, 'utf-8')
from urllib.parse import unquote
text = unquote(text, 'utf-8')
'''
###############################################################################
'''获取cookie，待验证，获取不全？
import http.cookiejar, urllib.request  # 导入要用到的cookiejar，request

url = "https://www.zhipin.com/gongsir/ee6a1eb8c53f8a020nJ729q6.html"
coikie = http.cookiejar.CookieJar()  # 声明CookieJar这个对象
hanlder = urllib.request.HTTPCookieProcessor(coikie)  # 利用HTTPCookieProcessor创建handler
opener = urllib.request.build_opener(hanlder)  # 最后用build_opener构建opener
response = opener.open(url)  # 用opener中的open来打开链接，跟urlopen类似效果
for item in coikie:
    print(item.expires)
    print(time.time())
    print(item.name + "=" + item.value)
'''
###############################################################################
'''单个字体拆分分析轮廓坐标等
from fontTools.ttLib import TTFont
from fontTools.pens.freetypePen import FreeTypePen
from fontTools.misc.transform import Offset
import matplotlib.pyplot as plt

font = TTFont('1.woff')
# font.saveXML('1.xml')
uni_list = font.getGlyphOrder()[1:]
print(uni_list)
name = "uniED6F"
glyph = font.getGlyphSet()[name]  # 获取_TTGlyph实例
zuobiao = font['glyf'][name].coordinates  # 坐标
x = [i[0] for i in zuobiao]
y = [i[1] for i in zuobiao]
plt.scatter(x, y)
plt.show()
print(glyph._glyph.coordinates)  # 坐标
print(glyph._glyph.endPtsOfContours)  # 轮廓结束点
print("".join(list(map(str, list(glyph._glyph.flags)))))  # 点类型flag
print("".join(map(str, glyph._glyph.flags)))  # 点类型flag
pen = FreeTypePen(None)  # 实例化Pen子类
glyph.draw(pen)  # “画”出字形轮廓
pen.show()
# width, ascender, descender = glyph.width, font['OS/2'].usWinAscent, -font['OS/2'].usWinDescent # 获取字形的宽度和上沿以及下沿
# height = ascender - descender # 利用上沿和下沿计算字形高度
# pen.show(width=width, height=height, transform=Offset(0, -descender)) # 显示以及矫正
'''
###############################################################################
'''JS回传字符串转字典
t = '{"bookId": "gpwx7lb53d", "chapterId": "29yeh3o9kmb", "fields": "read", "sourceType": -1}'
t =  eval('(' + t + ')')
'''
###############################################################################
'''放大缩小图片
path1 = 'apple1.jpg'
path2 = 'apple2.jpg'
times = 0.25   # 图片放大倍数
with open("apple1.jpg", "wb") as f:
    f.write(data)
img = cv2.imread(path1)
x, y = img.shape[0:2]
size = (int(y * times), int(x * times))  # 注意排序是反的
img = cv2.resize(img, size)
cv2.imwrite(path2, img)
'''
###############################################################################
'''网页转word文档
import pypandoc     # 安装注意事项参考 https://www.cnblogs.com/alex-13/p/16079221.html

url = "http://www.ccgp-jiangxi.gov.cn/web/jyxx/002006/002006007/20221109/7ff60394-578a-4ac3-b364-6900af6d07d7.html"
resp = requests.get(url, headers=headers)
html = resp.text
# output = pypandoc.convert_file('login.html', 'docx', outputfile="file1.docx")  # 将网页直接转换成docx
output = pypandoc.convert_text(html, 'docx', 'html', outputfile="file1.docx")  # 将 html 代码转化成docx
'''
###############################################################################
'''word内嵌表格转字典
# 首先要pip install python-docx
# 如果原文件是doc格式，那就先转成docx
from docx import Document
import pandas as pd

path = "./file1.docx"
docx = Document(path)
table_s = docx.tables  # 返回一个Table对象的列表
list_ = []  # 初始化一个空列表，用来装后面的dict_
for table in table_s:  # 循环所有的表格列表
    dict_ = {}
    for y in range(len(table.columns)):     # 适用于首行标题索引（即按列建立字典）
        key = table.cell(0, y).text
        li = []
        for x in range(1, len(table.rows)):
            value = table.cell(x, y).text
            li.append(value)
        dict_[key] = li
    print(dict_)
    dict_ = {}
    for y in range(len(table.rows)):     # 适用于首列标题索引（即按行建立字典）
        key = table.cell(y, 0).text
        li = []
        for x in range(1, len(table.columns)):
            value = table.cell(y, x).text
            li.append(value)
        dict_[key] = li
    print(dict_)
'''
###############################################################################
'''创建临时文件
import tempfile

fp = tempfile.TemporaryFile()
# fp = tempfile.TemporaryFile(mode='w+', dir=r'C:\Y\Case', delete=False)  # 默认以二进制写入，指定mode后可以字符串写入，不用encode、decode
print(fp.name)
fp.write(resp.content)
# 将文件指针移到开始处，准备读取文件
fp.seek(0)
print(fp.read().decode()) # 输出刚才写入的内容
# 关闭文件，该文件将会被自动删除
fp.close()  # 貌似不用clos，会自动删除
'''
###############################################################################
'''word转txt
path = 'file1.docx'
import docx

file = docx.Document(path)
print("段落数:" + str(len(file.paragraphs)))
for para in file.paragraphs:
    print(para.text)
for i in range(len(file.paragraphs)):
    print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)
'''
###############################################################################
'''word自动化，分段操作（文章段落、表格、图片）
path = 'file1.docx'
import docx

file = docx.Document(path)
print("段落数:" + str(len(file.paragraphs)))
for paragraph in file.paragraphs:
    for run in paragraph.runs:
        if "公共资源" in run.text:
            run.text = "江西xxxx交易网"
# file.paragraphs[1].runs[1].text = "name"
file.paragraphs[1].text = "name"
file.save('test.docx')
from docx import Document  # 导入模块

document = Document("test.docx")  # 设置打开的模板
for table in document.tables:  # 遍历所有表格
    print(len(table.rows))  # 打印当前表格行数
    print(len(table.columns))  # 打印当前表格列数
    for row_index, row in enumerate(table.rows):
        print([cell.text for cell in row.cells])
        for col_index, cell in enumerate(row.cells):
            if "A032099" in cell.text:
                cell.text = cell.text.replace("A032099", "123456")
        print([cell.text for cell in row.cells])
document.save("test2.docx")
'''
###############################################################################
'''
# 当前时间时间戳
# now = 1669878660
now = int(time.time())
#转换为其他日期格式,如:"%Y-%m-%d %H:%M:%S"
timeArray = time.localtime(now)
otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
print(otherStyleTime)
print(time.time())
print(time.asctime())
print(time.localtime())
'''
###############################################################################
'''调用带参数返回值的字符串函数
func_str = """
def func(m, n):
    return m * n
"""
namespace = {}
fun = compile(func_str, '<string>', 'exec')
exec(fun, namespace)
ret = namespace['func'](3, 4)
print(ret)
'''
###############################################################################
''' 装饰器带与不带参数
import json

# deco_func 与 indent不会同时存在
def json_output(deco_func=None, indent=None):
    print("deco_func=%s , indent=%s" % (deco_func, indent))

    def actual_decorator(func):
        print('func=%s' % func)
        print('......')
        def inner(*args, **kwargs):
            result = func(*args, **kwargs)
            print(result, 3)
            return json.dumps(result, indent=indent)
        return inner
    if deco_func:
        print(deco_func, 1)
        return actual_decorator(deco_func)  # 返回执行后的值（函数为参数时，直接返回二层装饰器的真实调用）
    else:
        print(actual_decorator, 2)
        return actual_decorator  # 返回函数（三层装饰器本质：返回代入参数以后的可调用的实际二层装饰器方法）

@json_output
def A():
    return {'status': '111'}

@json_output()
def B():
    return {'status': '222'}

@json_output(indent=8)
def C():
    return {'status': '333'}

# print(A())
# print(B())
# print(C())
'''
###############################################################################
''' 列表推导式 转 生成器
a = [i for i in range(4)]
print(*a)
a = iter(range(4))
print(*a)
'''
###############################################################################
''' 嵌套列表展开
li = [[1, 2], [3, 4], [5, 6],[[7]]]

print([item for sublist in li for item in sublist])     # 有缺陷
print(sum(li, []))      # 有缺陷

def flatten(li):    # 采用递归完美
    return sum(([x] if not isinstance(x, list) else flatten(x) for x in li), [])
print(flatten(li))
'''
###############################################################################
''' 列表添加指定索引起始值并还原，以推导式输出
li = [1,3,5,7]
index = 11
ll = [(i, j) for i, j in enumerate(li, index)]
print(ll)

index = 'a'
ll = [[chr(i), j] for i, j in enumerate(li, ord(index))]
print(ll)

l = [item[1] for item in ll]
print(l)
'''
###############################################################################
''' 元祖列表强行扩充其列表元素
t = ([1, 2, 3],)
t[0].extend([4])
print(t)
'''
###############################################################################
''' 多层嵌套装饰 ---> 外层输出：先上后下，内层返回输出：先下后上输出，最终返回：从下而上由内向外包装
# 方法一：(推荐)
import sys
def public_(name, params, func):
    def inner(*args, **kwargs):
        result = f"<{params}>{func()}</{params}>"
        print(name, '--->', result)
        return result
    return inner

def add_1(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        return public_(name, params, func)
    return wrapper

def add_2(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        return public_(name, params, func)
    return wrapper
def add_3(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        return public_(name, params, func)
    return wrapper

# 方法二：
def add_a(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        def inner(*args, **kwargs):
            result = f"<{params}>" + func() + f"</{params}>"
            print(add_a.__name__, params, result)
            return result
        return inner
    return wrapper

def add_b(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        def inner(*args, **kwargs):
            result = f"<{params}>{func()}</{params}>"
            print(add_b.__name__, params, result)
            return result
        return inner
    return wrapper
def add_c(params):
    name = sys._getframe().f_code.co_name
    print(name)
    def wrapper(func):
        def inner(*args, **kwargs):
            result = f"<{params}>{func()}</{params}>"
            print(add_c.__name__, params, result)
            return result
        return inner
    return wrapper

@add_3('3')
@add_2('2')
@add_1('1')
@add_c('c')
@add_b('b')
@add_a('a')
def content():
    sss = "人生苦短"
    print(sss, 'start')
    return sss

result = content()
print(result, 'end')
'''
###############################################################################
''' 获取该函数名称
# 方法一：（函数外部）
def fun_name():
    pass
z = fun_name
print(z.__name__)
print(getattr(z,'__name__'))

# 方法二：（函数内部）
import inspect
def fun_name():
    print(inspect.stack()[0][3])
fun_name()

# 方法三：（函数内部）
def fun_name():
    import sys
    print(sys._getframe().f_code.co_name)
fun_name()

# 方法四：（通过装饰器获取该函数名称）
def dec_name(f):
    name = f.__name__
    def new_f(*a, **ka):
        return f(*a, __name__=name, **ka)
    return new_f

@dec_name
def fun_name(__name__):
    print(__name__)

fun_name()
'''
###############################################################################
'''tkinter窗口递归叠加打开，注意按钮command一般不带参数不带括号，若需要得用闭包嵌套函数
x = 0
def get_click(window):
    def run():
        global x
        x += 1
        root = tkinter.Toplevel(window)     # 重点
        btn1 = tkinter.Button(root, text=f'按钮{x}', command=get_click(root))
        btn1.place(x=80, y=80)

    return run

root = tkinter.Tk()
btn = tkinter.Button(root, text='按钮0', command=get_click(root))
btn.place(x=80, y=80)
root.mainloop()
'''
###############################################################################
'''mysql windows版安装使用   网友参考：https://www.cnblogs.com/zhangkanghui/p/9613844.html
下载地址：https://dev.mysql.com/downloads/mysql/
（1）配置环境变量
变量名：MYSQL_HOME
变量值：C:\SecureBrowser360\mysql-8.0.31-winx64
cmd管理员方式打开进入 bin 目录
bin> mysqld --initialize-insecure --user=mysql
bin> mysqld -install
bin> net start mysql
登录MySQL
mysql> mysql -u root -p
查询用户密码
mysql> select host,user,authentication_string from mysql.user;
设置（或修改）root用户密码
mysql> use mysql
提别注意：下面这个修改密码的方式不正确，可能是因为版本问题。最近解决了。
mysql> update mysql.user set authentication_string=("root") where user="root"; 
解决方案如下：
mysql> ALTER USER 'root'@'localhost' IDENTIFIED WITH mysql_native_password BY 'root';
mysql> flush privileges;  
#作用：相当于保存，执行此命令后，设置才生效，若不执行，还是之前的密码不变
mysql> quit
'''
###############################################################################
''' 正则 “零宽断言” 用法
result = """
All User Profile : ziroom101
All User Profile : ziroom103_5G
All User Profile : ziroom103
All User Profile : Jan
"""
li = re.findall('(?<=\s:\s).*', result)
print(li)
li = re.findall('.*(?=\s:\s)', result)
print(li)
li = re.findall('(?:User)|(?:\s:\s)', result)
print(li)
'''
###############################################################################
''' 非侵入式地调试方式，通过装饰器形式更美观的展示调试运行过程
import pysnooper

# 方法一：针对函数用装饰器
@pysnooper.snoop()

# 方法二：针对某条代码，如调试 “海象运算符”
with pysnooper.snoop(output='Debug.log', prefix='前缀：'):     # 输出指定到文件 & 加前缀区分
    if a := 5 > 1:
        print(a, 'do sth!', a := 5)

    n = 3
    while (n := n - 1) + 1:
        print('do sth!', n)

    # fp = open("test.txt", "r")
    # while line := fp.readline():
    #     print(line.strip())
'''
###############################################################################
''' # 二值交换
a = 1
b = 2

c = a
a = b
b = c
print(a, b)
c = id(a)
a = ctypes.cast(id(b), ctypes.py_object).value
b = ctypes.cast(c, ctypes.py_object).value
print(a, b)
a = a + b
b = a - b
a = a - b
print(a, b)
a = a ^ b
b = a ^ b
a = a ^ b
print(a, b)
'''
###############################################################################
''' 查看python函数源代码、路径、支持参数等
import json
import inspect
import dill     # 第三方库
# source = inspect.getsource(json.dumps)
# source = inspect.getfile(json.dumps)
# source = inspect.signature(json.dumps)
source = dill.source.getsource(json.dumps)
print(source)
'''
###############################################################################
''' 将Python代码封装成库，实质是将代码目录拷贝到 C:\Python39\Lib\site-packages\ 目录，例如：/xujian/ 目录下代码 打包成 .whl
1、将代码全部放到 ./xujian/ 目录下
2、新建 ./setup.py 内容如下：
from setuptools import setup
packages = ['xujian']  # 需要封装的包的目录列表
setup(
    name='xujian',  # 打包后的包名，即安装包名
    version='1.0',  # 版本
    description='this is a program for xxx',  # 描述
    author='xu jian',  # 作者
    author_email='',  # 作者邮箱
    packages=packages,
)
# 3、打包 python39 ./setup.py sdist bdist_wheel  # 会多一个.zip包
3、打包 python39 ./setup.py bdist_wheel
4、安装 python39 -m pip install ./dist/xujian-1.0-py3-none-any.whl
5、使用 from xujian.test import fn
'''
###############################################################################
''' 生成临时身份令牌，带过期时间签名，貌似用处不大
import itsdangerous

salt = 'sdaf'  # 加盐，加调料，即密钥
t = itsdangerous.TimedJSONWebSignatureSerializer(salt, expires_in=600)  # 过期时间600秒
info = {'username': 'yangfan', 'user_id': 1}  # 加密的数据

# =========加密token============
token = t.dumps(info).decode()  # 指定编码格式
print(token)  # 生成以 eyJ 开头加密数据

print(token)
# =========解密token============
res = t.loads(token)
print(t.secret_key)
print(res)
# 当超时或值有误则会报错
'''
###############################################################################
''' 单例模式，主要功能是减少多次实例化调用运行内存(通过类私有变量赋值，下次调用被覆盖)和特殊方法在初始化时只运行一次的情况
class Solution:
    # 定义类变量
    # 记录第一个被创建对象的引用，代表着类的私有属性
    _instance = None
    # 记录是否执行过初始化动作
    init_flag = False

    def __init__(self, name, data):
        self.name = name
        self.data = data
        # 使用类名调用类变量,不能直接访问。
        if Solution.init_flag:
            return
        self.xml_load(self.data)
        # 修改类属性的标记
        Solution.init_flag = True

    def __new__(cls, *args, **kwargs):
        # 判断该类的属性是否为空；对第一个对象没有被创建，我们应该调用父类的方法，为第一个对象分配空间
        if cls._instance == None:
            # 把类属性中保存的对象引用返回给python的解释器
            cls._instance = object.__new__(cls)
        # 如果cls._instance不为None,直接返回已经实例化了的实例对象
        return cls._instance

    def xml_load(self, data):
        print("初始化init", self.name, data)

    def Parser(self):
        print("解析完成finish", self.name)


a = Solution("A11", 10)  # 第一次实例化对象地址，后面创建都是在该地址上进行的
a.Parser()
b = Solution("A12", 20)  # b把a覆盖掉
b.Parser()
print(id(a))
print(id(b))
print(a.name)
print(b.name)

xxx = 123456

import _ctypes

# 通过_ctypes的api进行对内存地址寻找对象
obj = _ctypes.PyObj_FromPtr(id(xxx))
# 打印出来通过内存地址寻找到的对象
print(obj)
obj = _ctypes.PyObj_FromPtr(id(b))
obj.Parser()
'''
###############################################################################
''' JWT 签名，pip install pyjwt
import jwt
from datetime import datetime, timedelta

key = """123456"""

# 生成 JWT Token
def encode_token(key):
    payload = {
        'exp': datetime.now() + timedelta(minutes=30),  # 令牌过期时间
        'username': 'BigFish'  # 想要传递的信息,如用户名ID
    }
    encoded_jwt = jwt.encode(payload, key, algorithm='HS256')
    return encoded_jwt

# 解码 JWT Token
def decode_token(encoded_jwt, key):
    res = jwt.decode(encoded_jwt, key, algorithms='HS256', options={"verify_signature": False})
    return res
    
# 
header = base64.b64encode(json.dumps({"typ": "JWT", "alg": "HS256"}).encode()).decode().replace('=', '')
encoded_jwt = header + '.' + base64.b64encode(json.dumps(payload).encode()).decode().replace('=', '')
'''
###############################################################################
'''文本指定行号行读取
import linecache
file_path = r'C:\Y\Case\subcase\疫情大数据\ip_detail.txt'
line_number = random.randint(1,4)
def get_line_context(file_path, line_number):
    return linecache.getline(file_path, line_number).strip()
'''
###############################################################################
'''JavaScript常用的Hook脚本
https://www.cnblogs.com/xiaoweigege/p/14954648.html#evalfunction
'''
###############################################################################
'''TrWebOCR-开源的离线OCR
1、Windows部署Docker，下载地址：https://www.docker.com/products/docker-desktop/
2、配置参考：https://blog.csdn.net/paycho/article/details/127574796
3、需要重启系统
4、进入项目文件夹执行如下两行命令
# 如下命令大约需要执行700秒
C:\Y\Case\验证码\TrWebOCR-master>docker build -t trwebocr:latest .
C:\Y\Case\验证码\TrWebOCR-master>docker run -itd --rm -p 8089:8089 --name trwebocr trwebocr:latest

# 访问：http://localhost:8089/
'''
###############################################################################
'''时间戳转换
now = time.time()
print(now)

# 时间戳 => 结构化时间 (struct_time) => 时间字符串
time_str = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(now))
print(time_str)

# 时间字符串 => 结构化时间 (struct_time) => 时间戳
# time_str = "2022-03-18 10:54:00"
time_stamp = time.mktime(time.strptime(time_str, '%Y-%m-%d %H:%M:%S'))
print(time_stamp)
'''
###############################################################################
'''
# 将字典转成URL请求参数
from urllib.parse import urlencode
params = {'a': 1, 'b': 2}
result = urlencode(params)
print(result)

# 将URL请求参数转成字典
from urllib.parse import parse_qs, urlparse
url = 'https://so.csdn.net/so/search?spm=1000.2115.3001.4498&q=test&t=&u='
query = urlparse(url).query
params = parse_qs(query)
print(params)
'''
###############################################################################
'''
# 方法一：
from base64 import b64encode, b64decode

# hex -> base64
s = 'cafebabe'
b64 = b64encode(bytes.fromhex(s)).decode()
print('cafebabe in base64:', b64)

# base64 -> hex
s2 = b64decode(b64.encode()).hex()
print('yv66vg== in hex is:', s2)
assert s == s2


# 方法二：
import base64
import codecs

# Hex转化成 base64 格式
hex_text_dome = "c2889d8f"
base64_text = codecs.encode(codecs.decode(hex_text_dome, 'hex'), 'base64').decode()
print(base64_text)

# base64转化成Hex格式
base64_text_dome = "woidjw=="
hex_text = base64.b64decode(base64_text_dome).hex()
print(hex_text)
'''
###############################################################################
'''多个分隔符分割字符串
import re

text = "小明，起床。吃早饭；然后，去上学 123,2，3。4：5；6|7、8？9=10+11-13*14 15,17.18:19;20/21"
text_list = re.split("[，。：；|、？=+\-\*\s,.:;/]", text)
print(text_list)
text_list = re.split("，|。|；|\s", text)
print(text_list)
'''
###############################################################################
'''sqlite3
C:\Users\11390>sqlite3
sqlite> .open 'C:\Users\11390\AppData\Local\Google\Chrome\User Data\Default\Network\Cookies'
sqlite> .databases
main: C:\Users\11390\AppData\Local\Google\Chrome\User Data\Default\Network\Cookies r/w
sqlite> .tables
cookies  meta
sqlite> .schema cookies
CREATE TABLE cookies(creation_utc INTEGER NOT NULL,host_key TEXT NOT NULL,top_frame_site_key TEXT NOT NULL,name TEXT NOT NULL,value TEXT NOT NULL,encrypted_value BLOB NOT NULL,path TEXT NOT NULL,expires_utc INTEGER NOT NULL,is_secure INTEGER NOT NULL,is_httponly INTEGER NOT NULL,last_access_utc INTEGER NOT NULL,has_expires INTEGER NOT NULL,is_persistent INTEGER NOT NULL,priority INTEGER NOT NULL,samesite INTEGER NOT NULL,source_scheme INTEGER NOT NULL,source_port INTEGER NOT NULL,is_same_party INTEGER NOT NULL);
CREATE UNIQUE INDEX cookies_unique_index ON cookies(host_key, top_frame_site_key, name, path);
sqlite>
sqlite> select * from cookies where name='cf_clearance';
13324879501376613|.800xsw.net||cf_clearance||v10       4??/bA?jR?o????{p|/|13356415725882825|1|1|13324879725882825|1|1|1|0|2|443|0
sqlite>
'''
###############################################################################
''' cookie 本地存放位置
windows:
C:\Users\11390\AppData\Local\Google\Chrome\User Data\Default\Network\Cookies
mac:
~/Library/Application Support/Google/Chrome/Default/Cookies
'''
###############################################################################
'''

'''
###############################################################################
'''

'''
###############################################################################
'''

'''
###############################################################################
'''

'''